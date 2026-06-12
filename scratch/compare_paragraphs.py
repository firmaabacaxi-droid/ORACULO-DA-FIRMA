import docx
import difflib

v2_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v2 lipe.docx"
v4_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v4.docx"

doc2 = docx.Document(v2_path)
doc4 = docx.Document(v4_path)

p2 = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
p4 = [p.text.strip() for p in doc4.paragraphs if p.text.strip()]

print(f"v2 paragraphs count (non-empty): {len(p2)}")
print(f"v4 paragraphs count (non-empty): {len(p4)}")

diff = list(difflib.unified_diff(p2, p4, fromfile="v2_lipe", tofile="v4", lineterm=""))
if not diff:
    print("No differences in paragraphs text between V2 Lipe and V4!")
else:
    print("Differences found between V2 Lipe and V4:")
    for line in diff:
        print(line)
