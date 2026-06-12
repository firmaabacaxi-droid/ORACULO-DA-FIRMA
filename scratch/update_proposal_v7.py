import docx
from docx.shared import Pt

def update_cell_text_preserve_format(cell, text):
    p = cell.paragraphs[0]
    if not p.runs:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""

doc_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v6.docx"
out_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v7.docx"

doc = docx.Document(doc_path)

print("Processando tabelas da Proposta...")

# 1. Tabela 1 (ENAP Videos) - doc.tables[1]
t1 = doc.tables[1]
update_cell_text_preserve_format(t1.rows[1].cells[1], "R$ 1.500,00") # valor unitário
update_cell_text_preserve_format(t1.rows[1].cells[3], "R$ 15.000,00") # valor total
update_cell_text_preserve_format(t1.rows[2].cells[3], "R$ 15.000,00") # subtotal
update_cell_text_preserve_format(t1.rows[3].cells[3], "R$ 1.092,00") # NF 7,28%
update_cell_text_preserve_format(t1.rows[4].cells[3], "R$ 16.092,00") # total

# 2. Tabela 2 (Resumo Consolidado) - doc.tables[2]
t2 = doc.tables[2]
# Escopo 1 - Corrigir subtotal e imposto para bater com Tabela 0
update_cell_text_preserve_format(t2.rows[1].cells[1], "R$ 46.630,00") # subtotal
update_cell_text_preserve_format(t2.rows[1].cells[2], "R$ 3.395,00") # imposto
update_cell_text_preserve_format(t2.rows[1].cells[3], "R$ 50.025,00") # total

# Escopo 2 - Atualizar com os novos valores
update_cell_text_preserve_format(t2.rows[2].cells[1], "R$ 15.000,00") # subtotal
update_cell_text_preserve_format(t2.rows[2].cells[2], "R$ 1.092,00") # imposto
update_cell_text_preserve_format(t2.rows[2].cells[3], "R$ 16.092,00") # total

# Total Consolidado
update_cell_text_preserve_format(t2.rows[3].cells[3], "R$ 66.117,00")

print("[OK] Tabelas atualizadas.")

print("Processando parágrafos de Condições de Pagamento...")

# 3. Atualizar parágrafos de pagamento
for p in doc.paragraphs:
    if "Escopo 1" in p.text and "Documentário Principal" in p.text and "faturado mediante" in p.text:
        # Atualizar milestones do documentário
        p.text = "• Escopo 1 — Documentário Principal: 50% do valor do documentário (R$ 25.012,50) faturado mediante a entrega da primeira versão do roteiro documental, e 50% (R$ 25.012,50) faturado mediante a entrega do produto final (arquivo master aprovado)."
        print("[OK] Parágrafo pagamento Escopo 1 atualizado.")
    elif "Escopo 2" in p.text and "Vídeos Instrucionais" in p.text and "faturado mediante" in p.text:
        # Atualizar milestones dos vídeos
        p.text = "• Escopo 2 — Vídeos Instrucionais (Adicional): 50% do valor das videoaulas (R$ 8.046,00) faturado mediante a entrega do primeiro vídeo instrucional finalizado, e 50% (R$ 8.046,00) faturado mediante a entrega do décimo (último) vídeo instrucional finalizado."
        print("[OK] Parágrafo pagamento Escopo 2 atualizado.")

# Salvar documento
doc.save(out_path)
print(f"[OK] Documento v7 salvo com sucesso em: {out_path}")
