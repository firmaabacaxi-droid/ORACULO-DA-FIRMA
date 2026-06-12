import docx

versions = {
    "v2_lipe": r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v2 lipe.docx",
    "v3": r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v3.docx",
    "v4": r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v4.docx",
    "v5": r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v5.docx",
    "v6": r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v6.docx",
}

for name, path in versions.items():
    print(f"\n=================== {name.upper()} ===================")
    try:
        doc = docx.Document(path)
        print(f"Number of paragraphs: {len(doc.paragraphs)}")
        print(f"Number of tables: {len(doc.tables)}")
        
        # Print tables
        for t_idx, table in enumerate(doc.tables):
            print(f"\nTabela {t_idx}:")
            for r_idx, row in enumerate(table.rows):
                # unique cells check
                cells_text = []
                for cell in row.cells:
                    cells_text.append(cell.text.strip().replace('\n', ' '))
                print(f"  Linha {r_idx}: {cells_text}")
                
        # Find if there are payment paragraphs
        print("\nPayment/Validade paragraphs:")
        for idx, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if "Pagamento" in text or "Validade" in text or "Cronograma" in text:
                print(f"  {idx}: {text}")
    except Exception as e:
        print(f"Error reading {name}: {e}")
