import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Emu

def copy_paragraph_format(src, dst):
    dst.paragraph_format.left_indent = src.paragraph_format.left_indent
    dst.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
    dst.paragraph_format.space_before = src.paragraph_format.space_before
    dst.paragraph_format.space_after = src.paragraph_format.space_after
    dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing
    dst.paragraph_format.alignment = src.paragraph_format.alignment

def copy_run_format(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    dst.font.name = src.font.name
    dst.font.size = src.font.size
    dst.font.color.rgb = src.font.color.rgb

def format_new_cell(cell, text, align, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.bold = bold

doc_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v4.docx"
out_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v5.docx"

doc = docx.Document(doc_path)

print("Processando paragrafos...")

# 1. Inserir bullet em "O que esta incluido:"
idx_exclusao = -1
for i, p in enumerate(doc.paragraphs):
    if "O que não está incluído:" in p.text:
        idx_exclusao = i
        break

if idx_exclusao != -1:
    p_exclusao = doc.paragraphs[idx_exclusao]
    p_ref = doc.paragraphs[idx_exclusao - 1] # O ultimo bullet do "O que esta incluido:"
    
    new_p = p_exclusao.insert_paragraph_before()
    copy_paragraph_format(p_ref, new_p)
    
    # Adicionar runs
    r_bullet = new_p.add_run(chr(183) + "  ")
    copy_run_format(p_ref.runs[0], r_bullet)
    
    r_text = new_p.add_run("Pós-produção de vídeos instrucionais: edição e finalização de 10 (dez) videoaulas e tutoriais de tela sobre a plataforma Brasil Participativo, destinados à Escola Virtual de Governo (ENAP).")
    copy_run_format(p_ref.runs[1], r_text)
    print("[OK] Bullet de inclusao inserido.")
else:
    print("[ERROR] Nao encontrou 'O que nao esta incluido:'")

# 2. Inserir Diretrizes Tecnicas antes de "Cronograma"
idx_cronograma = -1
for i, p in enumerate(doc.paragraphs):
    if "Cronograma" in p.text and p.text.strip().startswith(chr(8226)):
        idx_cronograma = i
        break

if idx_cronograma != -1:
    p_cronograma = doc.paragraphs[idx_cronograma]
    # Usaremos o estilo do cabecalho "Nossa Proposta" ou "Cronograma" como referencia para o novo cabecalho
    p_ref_head = p_cronograma
    p_ref_body = doc.paragraphs[idx_cronograma - 1] # paragrafo normal anterior
    
    # 2.1 Inserir Heading das Diretrizes
    p_head = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_head, p_head)
    # Adiciona espacamento extra antes do heading para separar da secao anterior
    p_head.paragraph_format.space_before = Pt(24)
    r_bullet_head = p_head.add_run(chr(8226) + "  ")
    copy_run_format(p_ref_head.runs[0], r_bullet_head)
    r_text_head = p_head.add_run("Diretrizes Técnicas (Vídeos Instrucionais)")
    copy_run_format(p_ref_head.runs[1], r_text_head)
    
    # 2.2 Inserir Corpo 1
    p_body1 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_body1)
    p_body1.paragraph_format.space_after = Pt(12)
    r_body1 = p_body1.add_run("Os 10 (dez) vídeos instrucionais adicionais serão gravados diretamente pela equipe do cliente/Presidência e fornecidos à Firma Abacaxi para pós-produção. Cada vídeo terá duração estimada entre 5 e 10 minutos (totalizando de 1h a 2h de material finalizado), distribuídos entre os formatos de tutorial de tela (captura de tela com locução) e videoaula (exposição do instrutor + compartilhamento de tela).")
    copy_run_format(p_ref_body.runs[0], r_body1)
    
    # 2.3 Inserir Corpo 2
    p_body2 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_body2)
    p_body2.paragraph_format.space_after = Pt(24) # mais espacamento depois do ultimo paragrafo
    r_body2 = p_body2.add_run("Como diretriz crítica para a viabilidade e cronograma do projeto, o material bruto de gravação deve ser fornecido em conformidade técnica com as orientações de captação da produtora (áudio limpo e sem ruídos, enquadramento de câmera estável e adequado, e capturas de tela em alta resolução). Desvios significativos que exijam refilmagem ou correções complexas de sincronização de áudio e tela demandarão novos prazos e orçamento adicional a ser acordado via aditivo contratual.")
    copy_run_format(p_ref_body.runs[0], r_body2)
    
    print("[OK] Secao de Diretrizes Tecnicas inserida.")
else:
    print("[ERROR] Nao encontrou heading 'Cronograma'")

# 3. Inserir bullet em "Cronograma"
idx_feedback = -1
for i, p in enumerate(doc.paragraphs):
    if "O prazo de feedback do LabLivre" in p.text:
        idx_feedback = i
        break

if idx_feedback != -1:
    p_feedback = doc.paragraphs[idx_feedback]
    p_ref_crono = doc.paragraphs[idx_feedback - 1] # O ultimo bullet do Cronograma
    
    new_p_crono = p_feedback.insert_paragraph_before()
    copy_paragraph_format(p_ref_crono, new_p_crono)
    
    r_bullet_crono = new_p_crono.add_run(chr(183) + "  ")
    copy_run_format(p_ref_crono.runs[0], r_bullet_crono)
    
    r_text_crono = new_p_crono.add_run("Período de execução dos vídeos instrucionais: de julho a novembro de 2026 (desenvolvido em fluxo contínuo paralelo, acompanhando as demandas de capacitação da plataforma).")
    copy_run_format(p_ref_crono.runs[1], r_text_crono)
    
    print("[OK] Bullet de cronograma inserido.")
else:
    print("[ERROR] Nao encontrou paragrafo de feedback do LabLivre")


# 4. Modificar Tabela de Orçamento
print("Modificando tabela...")
table = doc.tables[0]

# Inserir nova linha no indice 6 (antes do Subtotal)
row = table.add_row()
table._tbl.remove(row._tr)
table._tbl.insert(8, row._tr) # XML index = row index + 2 = 6 + 2 = 8

format_new_cell(row.cells[0], "Edição e finalização de vídeos instrucionais (ENAP)", WD_ALIGN_PARAGRAPH.LEFT)
format_new_cell(row.cells[1], "R$ 932,20", WD_ALIGN_PARAGRAPH.CENTER)
format_new_cell(row.cells[2], "10 vídeos", WD_ALIGN_PARAGRAPH.CENTER)
format_new_cell(row.cells[3], "R$ 9.322", WD_ALIGN_PARAGRAPH.RIGHT)

# Atualizar as linhas de totais
# Agora as linhas de total mudaram de indice devido a insercao:
# Indice 7: Subtotal
# Indice 8: NF
# Indice 9: TOTAL
table.rows[7].cells[3].paragraphs[0].runs[0].text = "R$ 55.934"
table.rows[8].cells[3].paragraphs[0].runs[0].text = "R$ 4.066"
table.rows[9].cells[3].paragraphs[0].runs[0].text = "R$ 60.000"

print("[OK] Tabela de orcamento atualizada.")

# Salvar documento
doc.save(out_path)
print(f"[OK] Documento salvo com sucesso em: {out_path}")
