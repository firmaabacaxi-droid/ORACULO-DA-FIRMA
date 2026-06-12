import docx

doc_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v4.docx"
doc = docx.Document(doc_path)

print("=== PARAGRAFOS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{i}: {para.text}")

print("\n=== TABELAS ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTabela {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"  Linha {r_idx}: {cells}")
