import docx

path = r"C:\Users\User\Meu Drive\ORACULO- FIRMA ABACAXI 2026\PROJETOS\2026\Brasil Participativo\01_PROPOSTA\BrasilParticipativo_proposta_v6.docx"

try:
    doc = docx.Document(path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    
    # Print tables
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTabela {t_idx}:")
        for r_idx, row in enumerate(table.rows):
            cells_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"  Linha {r_idx}: {cells_text}")
            
    # Print key paragraphs
    print("\nPayment/Validade/Cronograma paragraphs:")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if any(w in text.lower() for w in ["pagamento", "validade", "cronograma", "total", "faturado", "R$"]):
            print(f"  {idx}: {text}")
except Exception as e:
    print(f"Error reading doc: {e}")
