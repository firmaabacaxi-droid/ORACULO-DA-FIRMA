import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import copy

def copy_paragraph_format(src, dst):
    dst.paragraph_format.left_indent = src.paragraph_format.left_indent
    dst.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
    dst.paragraph_format.space_before = src.paragraph_format.space_before
    dst.paragraph_format.space_after = src.paragraph_format.space_after
    dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing
    dst.paragraph_format.alignment = src.paragraph_format.alignment

def copy_run_format(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    dst.font.name = src.font.name
    dst.font.size = src.font.size
    dst.font.color.rgb = src.font.color.rgb

def update_cell_text_preserve_format(cell, text):
    p = cell.paragraphs[0]
    if not p.runs:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    else:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""

def create_table_from_rows(original_table, row_indices):
    tbl_clone = copy.deepcopy(original_table._tbl)
    tr_elements = tbl_clone.xpath('w:tr')
    for tr in tr_elements:
        tbl_clone.remove(tr)
        
    for idx in row_indices:
        tr_copy = copy.deepcopy(original_table.rows[idx]._tr)
        tbl_clone.append(tr_copy)
        
    return tbl_clone

doc_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v4.docx"
out_path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v6.docx"

doc = docx.Document(doc_path)

print("Processando paragrafos...")

# 1. Inserir Escopo Adicional antes do Cronograma e renomear Cronograma
idx_cronograma = -1
for i, p in enumerate(doc.paragraphs):
    if "Cronograma" in p.text and (p.text.strip().startswith(chr(8226)) or p.text.strip().startswith("•")):
        idx_cronograma = i
        break

if idx_cronograma != -1:
    p_cronograma = doc.paragraphs[idx_cronograma]
    p_ref_head = p_cronograma
    p_ref_body = doc.paragraphs[10] # parágrafo normal de corpo garantido com runs
    p_ref_bullet = doc.paragraphs[12] # bullet reference (usa EM DASH no original)
    
    # 1.1 Inserir Heading do Escopo Adicional
    p_head = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_head, p_head)
    p_head.paragraph_format.space_before = Pt(24)
    r_bullet_head = p_head.add_run(chr(8226) + "  ")
    copy_run_format(p_ref_head.runs[0], r_bullet_head)
    r_text_head = p_head.add_run("Escopo Adicional: Pós-Produção de Vídeos Instrucionais (ENAP)")
    copy_run_format(p_ref_head.runs[1], r_text_head)
    
    # 1.2 Inserir Corpo 1
    p_body1 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_body1)
    p_body1.paragraph_format.space_after = Pt(12)
    r_body1 = p_body1.add_run("Propomos a pós-produção (edição e finalização) de 10 (dez) vídeos de pequenas capacitações destinados à Escola Virtual de Governo (ENAP).")
    copy_run_format(p_ref_body.runs[0], r_body1)
    
    # 1.3 Inserir "O que está incluído..."
    p_incl_head = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_incl_head)
    p_incl_head.paragraph_format.space_after = Pt(12)
    r_incl_head = p_incl_head.add_run("O que está incluído neste escopo adicional:")
    copy_run_format(p_ref_body.runs[0], r_incl_head)
    r_incl_head.bold = True
    
    # 1.4 Bullet 1
    p_b1 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_bullet, p_b1)
    r_bullet1 = p_b1.add_run(chr(8212) + " ") # usa em-dash
    copy_run_format(p_ref_bullet.runs[0], r_bullet1)
    r_text1 = p_b1.add_run("Pós-produção completa: edição do material bruto, sincronização de áudio/tela, tratamento básico de imagem e áudio, inserção de cartelas, slides de apoio e vinhetas padrão de abertura e encerramento.")
    copy_run_format(p_ref_bullet.runs[1], r_text1)
    
    # 1.5 Bullet 2
    p_b2 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_bullet, p_b2)
    r_bullet2 = p_b2.add_run(chr(8212) + " ")
    copy_run_format(p_ref_bullet.runs[0], r_bullet2)
    r_text2 = p_b2.add_run("Entregáveis: 10 (dez) vídeos instrucionais finalizados, com duração estimada de 5 a 10 minutos cada (totalizando de 1h a 2h de material finalizado).")
    copy_run_format(p_ref_bullet.runs[1], r_text2)
    
    # 1.6 Inserir "O que não está incluído..."
    p_excl_head = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_excl_head)
    p_excl_head.paragraph_format.space_before = Pt(12)
    p_excl_head.paragraph_format.space_after = Pt(12)
    r_excl_head = p_excl_head.add_run("O que não está incluído neste escopo adicional:")
    copy_run_format(p_ref_body.runs[0], r_excl_head)
    r_excl_head.bold = True
    
    # 1.7 Bullet 3
    p_b3 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_bullet, p_b3)
    r_bullet3 = p_b3.add_run(chr(8212) + " ")
    copy_run_format(p_ref_bullet.runs[0], r_bullet3)
    r_text3 = p_b3.add_run("Captação de imagem, gravação de voz/locução ou gravação de tela (responsabilidade integral do cliente).")
    copy_run_format(p_ref_bullet.runs[1], r_text3)
    
    # 1.8 Bullet 4
    p_b4 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_bullet, p_b4)
    r_bullet4 = p_b4.add_run(chr(8212) + " ")
    copy_run_format(p_ref_bullet.runs[0], r_bullet4)
    r_text4 = p_b4.add_run("Desenvolvimento de roteiros ou produção de locução original (responsabilidade integral do cliente).")
    copy_run_format(p_ref_bullet.runs[1], r_text4)
    
    # 1.9 Bullet 5
    p_b5 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_bullet, p_b5)
    r_bullet5 = p_b5.add_run(chr(8212) + " ")
    copy_run_format(p_ref_bullet.runs[0], r_bullet5)
    r_text5 = p_b5.add_run("Correções complexas decorrentes de falhas na captação bruta original ou refilmagens.")
    copy_run_format(p_ref_bullet.runs[1], r_text5)
    
    # 1.10 Inserir "Diretrizes Técnicas e de Cronograma:"
    p_dir_head = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_dir_head)
    p_dir_head.paragraph_format.space_before = Pt(12)
    p_dir_head.paragraph_format.space_after = Pt(12)
    r_dir_head = p_dir_head.add_run("Diretrizes Técnicas e de Cronograma:")
    copy_run_format(p_ref_body.runs[0], r_dir_head)
    r_dir_head.bold = True
    
    # 1.11 Parágrafo de Diretrizes
    p_body2 = p_cronograma.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_body2)
    p_body2.paragraph_format.space_after = Pt(24)
    r_body2 = p_body2.add_run("Como diretriz crítica para a viabilidade deste escopo adicional, o material bruto de gravação deve ser fornecido em conformidade técnica com as orientações de captação da produtora (áudio limpo e sem ruídos, enquadramento de câmera estável, e capturas de tela em alta resolução). Desvios significativos demandarão novos prazos e orçamento adicional a ser acordado via aditivo contratual. A execução das edições ocorrerá em fluxo contínuo de julho a novembro de 2026, de acordo com as demandas de capacitação da plataforma.")
    copy_run_format(p_ref_body.runs[0], r_body2)

    # 1.12 Renomear o Cronograma para Cronograma do Documentário
    for run in p_cronograma.runs:
        if "Cronograma" in run.text:
            run.text = run.text.replace("Cronograma", "Cronograma do Documentário")

    print("[OK] Escopo adicional inserido e Cronograma renomeado.")
else:
    print("[ERROR] Nao encontrou heading 'Cronograma'")


# 2. Modificar Seção de Condições de Pagamento
p_pagamento = None
p_validade = None
for p in doc.paragraphs:
    if "pagamento:" in p.text.lower():
        p_pagamento = p
    elif "validade desta proposta:" in p.text.lower():
        p_validade = p

if p_pagamento and p_validade:
    p_ref_body = doc.paragraphs[10] # garantido com runs
    p_ref_bullet = doc.paragraphs[12] # bullet reference
    p_ref_head = p_cronograma
    
    # 2.1 Inserir Heading
    p_pay_head = p_validade.insert_paragraph_before()
    copy_paragraph_format(p_ref_head, p_pay_head)
    p_pay_head.paragraph_format.space_before = Pt(24)
    r_bullet_pay = p_pay_head.add_run(chr(8226) + "  ")
    copy_run_format(p_ref_head.runs[0], r_bullet_pay)
    r_text_pay = p_pay_head.add_run("Condições de Pagamento por Escopo")
    copy_run_format(p_ref_head.runs[1], r_text_pay)
    
    # 2.2 Inserir texto intro
    p_pay_intro = p_validade.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_pay_intro)
    p_pay_intro.paragraph_format.space_after = Pt(12)
    r_pay_intro = p_pay_intro.add_run("O faturamento dos serviços será desatrelado por escopo, conforme as entregas e marcos específicos de cada projeto:")
    copy_run_format(p_ref_body.runs[0], r_pay_intro)
    
    # 2.3 Bullet 1: Documentário
    p_pay_b1 = p_validade.insert_paragraph_before()
    copy_paragraph_format(p_ref_bullet, p_pay_b1)
    p_pay_b1.paragraph_format.space_after = Pt(6)
    r_pay_bullet1 = p_pay_b1.add_run(chr(8212) + " ")
    copy_run_format(p_ref_bullet.runs[0], r_pay_bullet1)
    r_pay_text1 = p_pay_b1.add_run("Escopo 1 — Documentário Principal: 50% do valor do documentário (R$ 25.000,00) faturado mediante a entrega da primeira versão do roteiro documental, e 50% (R$ 25.000,00) faturado mediante a entrega do produto final (arquivo master aprovado).")
    copy_run_format(p_ref_bullet.runs[1], r_pay_text1)
    
    # 2.4 Bullet 2: Vídeos Instrucionais
    p_pay_b2 = p_validade.insert_paragraph_before()
    copy_paragraph_format(p_ref_bullet, p_pay_b2)
    p_pay_b2.paragraph_format.space_after = Pt(18)
    r_pay_bullet2 = p_pay_b2.add_run(chr(8212) + " ")
    copy_run_format(p_ref_bullet.runs[0], r_pay_bullet2)
    r_pay_text2 = p_pay_b2.add_run("Escopo 2 — Vídeos Instrucionais (Adicional): 50% do valor das videoaulas (R$ 7.500,00) faturado mediante a entrega do primeiro vídeo instrucional finalizado, e 50% (R$ 7.500,00) faturado mediante a entrega do décimo (último) vídeo instrucional finalizado.")
    copy_run_format(p_ref_bullet.runs[1], r_pay_text2)
    
    # Deletar o parágrafo de pagamento original
    p_pagamento._p.getparent().remove(p_pagamento._p)
    print("[OK] Condicoes de pagamento atualizadas.")
else:
    print("[ERROR] Nao encontrou paragrafo de Pagamento ou Validade")


# 3. Criar Tabelas de Orçamento
print("Processando tabelas...")
table_doc = doc.tables[0] # Tabela original do documentário

# Inserir cabeçalho do Escopo 1 antes da tabela 1
p_escopo1_head = doc.add_paragraph()
copy_paragraph_format(p_ref_body, p_escopo1_head)
p_escopo1_head.paragraph_format.space_before = Pt(12)
p_escopo1_head.paragraph_format.space_after = Pt(6)
r_escopo1_head = p_escopo1_head.add_run("Orçamento do Escopo 1 — Documentário Principal")
copy_run_format(p_ref_body.runs[0], r_escopo1_head)
r_escopo1_head.bold = True
r_escopo1_head.font.size = Pt(11)

tbl_element = table_doc._tbl
tbl_element.getparent().insert(tbl_element.getparent().index(tbl_element), p_escopo1_head._p)
print("[OK] Tabela 1 identificada (Documentario) e titulo inserido antes dela")

# Encontrar a nota de rodapé
p_nota = None
idx_nota = -1
for i, p in enumerate(doc.paragraphs):
    if "* inclui dire" in p.text.lower() or "* a diária de" in p.text.lower():
        p_nota = p
        idx_nota = i
        break

if idx_nota != -1:
    p_next = doc.paragraphs[idx_nota + 1]
    
    # 3.1 Inserir Heading para Tabela 2
    p_head2 = p_next.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_head2)
    p_head2.paragraph_format.space_before = Pt(24)
    p_head2.paragraph_format.space_after = Pt(6)
    r_head2 = p_head2.add_run("Orçamento do Escopo 2 — Vídeos Instrucionais (Adicional)")
    copy_run_format(p_ref_body.runs[0], r_head2)
    r_head2.bold = True
    r_head2.font.size = Pt(11)
    
    # 3.2 Inserir Tabela 2
    # Queremos: Header (0), Data (1), Subtotal (6), NF (7), Total (8)
    tbl2_xml = create_table_from_rows(table_doc, [0, 1, 6, 7, 8])
    p_head2._p.getparent().insert(p_head2._p.getparent().index(p_head2._p) + 1, tbl2_xml)
    
    table2 = docx.table.Table(tbl2_xml, doc)
    
    # Modificar valores da Tabela 2
    update_cell_text_preserve_format(table2.rows[1].cells[0], "Edição e finalização de vídeos instrucionais (ENAP) — tutoriais de tela e videoaulas")
    update_cell_text_preserve_format(table2.rows[1].cells[1], "R$ 1.398,30")
    update_cell_text_preserve_format(table2.rows[1].cells[2], "10 vídeos")
    update_cell_text_preserve_format(table2.rows[1].cells[3], "R$ 13.983")
    
    update_cell_text_preserve_format(table2.rows[2].cells[3], "R$ 13.983,00")
    update_cell_text_preserve_format(table2.rows[3].cells[3], "R$ 1.017,00")
    update_cell_text_preserve_format(table2.rows[4].cells[0], "TOTAL VÍDEOS INSTRUCIONAIS (ADICIONAL)")
    update_cell_text_preserve_format(table2.rows[4].cells[3], "R$ 15.000,00")
    
    # 3.3 Inserir Heading para Tabela 3 (Resumo Consolidado)
    p_head3 = p_next.insert_paragraph_before()
    copy_paragraph_format(p_ref_body, p_head3)
    p_head3.paragraph_format.space_before = Pt(24)
    p_head3.paragraph_format.space_after = Pt(6)
    r_head3 = p_head3.add_run("Resumo Consolidado do Investimento")
    copy_run_format(p_ref_body.runs[0], r_head3)
    r_head3.bold = True
    r_head3.font.size = Pt(11)
    
    # 3.4 Inserir Tabela 3
    # Queremos: Header (0), Data (1), Data (1), Total (8)
    tbl3_xml = create_table_from_rows(table_doc, [0, 1, 1, 8])
    p_head3._p.getparent().insert(p_head3._p.getparent().index(p_head3._p) + 1, tbl3_xml)
    
    table3 = docx.table.Table(tbl3_xml, doc)
    
    # Modificar valores da Tabela 3
    # Header
    update_cell_text_preserve_format(table3.rows[0].cells[0], "Escopo / Projeto")
    update_cell_text_preserve_format(table3.rows[0].cells[1], "Subtotal de Serviços")
    update_cell_text_preserve_format(table3.rows[0].cells[2], "Imposto (NF 7,28%)")
    update_cell_text_preserve_format(table3.rows[0].cells[3], "Total Consolidado")
    
    # Row 1: Documentário
    update_cell_text_preserve_format(table3.rows[1].cells[0], "Escopo 1 — Documentário Principal (LabLivre / Brasil Participativo)")
    update_cell_text_preserve_format(table3.rows[1].cells[1], "R$ 46.612,00")
    update_cell_text_preserve_format(table3.rows[1].cells[2], "R$ 3.388,00")
    update_cell_text_preserve_format(table3.rows[1].cells[3], "R$ 50.000,00")
    
    # Row 2: Vídeos instrucionais
    update_cell_text_preserve_format(table3.rows[2].cells[0], "Escopo 2 — Vídeos Instrucionais (ENAP) [ADICIONAL]")
    update_cell_text_preserve_format(table3.rows[2].cells[1], "R$ 13.983,00")
    update_cell_text_preserve_format(table3.rows[2].cells[2], "R$ 1.017,00")
    update_cell_text_preserve_format(table3.rows[2].cells[3], "R$ 15.000,00")
    
    # Row 3: Total consolidado
    update_cell_text_preserve_format(table3.rows[3].cells[0], "TOTAL CONSOLIDADO DA PARCERIA")
    update_cell_text_preserve_format(table3.rows[3].cells[3], "R$ 65.000,00")
    
    print("[OK] Tabelas 2 e 3 geradas e atualizadas.")
else:
    print("[ERROR] Nao encontrou a nota de rodape da tabela 1")

# Salvar documento
doc.save(out_path)
print(f"[OK] Documento v6 salvo com sucesso em: {out_path}")
