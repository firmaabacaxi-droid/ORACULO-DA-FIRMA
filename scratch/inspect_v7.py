import docx

path = r"c:\Users\User\Documents\ORACULO - FIRMA ABACAXI\output\propostas\BrasilParticipativo_proposta_v7.docx"
doc = docx.Document(path)

print("=== TABELAS V7 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTabela {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        cells_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"  Linha {r_idx}: {cells_text}")
        
print("\n=== PAGAMENTOS V7 ===")
for idx, p in enumerate(doc.paragraphs):
    if "Escopo" in p.text and "faturado mediante" in p.text:
        print(f"  {idx}: {p.text}")
